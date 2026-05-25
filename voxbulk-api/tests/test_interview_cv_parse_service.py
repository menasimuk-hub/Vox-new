from app.services.interview_cv_parse_service import parse_docx_bytes, parse_pdf_bytes


def test_parse_docx_extracts_contacts():
    from docx import Document
    import io

    doc = Document()
    doc.add_paragraph("Jane Candidate")
    doc.add_paragraph("Email: jane@example.com")
    doc.add_paragraph("Phone: +44 7700 900123")
    doc.add_paragraph("Skills: Python, React, SQL")
    buf = io.BytesIO()
    doc.save(buf)
    parsed = parse_docx_bytes(buf.getvalue(), "jane.docx")
    assert parsed.name
    assert parsed.email == "jane@example.com"
    assert parsed.phone
    assert parsed.quality in {"good", "low_quality"}


def test_parse_pdf_low_quality_on_empty():
    import fitz

    doc = fitz.open()
    doc.new_page()
    buf = doc.write()
    doc.close()
    parsed = parse_pdf_bytes(buf, "empty.pdf")
    assert parsed.quality == "low_quality"
