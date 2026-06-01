import io
import uuid
import zipfile

import pytest
from docx import Document
from sqlalchemy.orm import Session

from app.models.organisation import Organisation
from app.models.service_order import ServiceOrder
from app.models.user import User
from app.services.interview_intake_service import (
    abandon_empty_interview_draft,
    create_new_interview_draft,
    intake_cv_files,
    intake_mixed_files,
    is_empty_interview_draft,
    purge_empty_interview_drafts,
)
from app.services.platform_catalog_service import ServiceOrderService


def _docx_bytes(name_line: str, *, phone: str = "+44 7700 900123", email: str = "alex@example.com") -> bytes:
    doc = Document()
    doc.add_paragraph(name_line)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _zip_with_docx(files: list[tuple[str, bytes]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for fname, data in files:
            zf.writestr(fname, data)
    return buf.getvalue()


@pytest.fixture()
def db_session():
    from app.core.database import Base, get_engine, get_sessionmaker
    import app.models  # noqa: F401

    engine = get_engine()
    Base.metadata.create_all(bind=engine)
    session = get_sessionmaker()()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture()
def interview_order(db_session: Session):
    org = Organisation(name="Test Org")
    user = User(email=f"intake-{uuid.uuid4().hex[:8]}@test.com", password_hash="x", is_active=True)
    db_session.add(org)
    db_session.add(user)
    db_session.flush()
    order = ServiceOrderService.create_order(
        db_session,
        org_id=org.id,
        user_id=user.id,
        service_code="interview",
        title="Interview draft",
        config={},
    )
    return order


def test_intake_cv_zip_creates_recipient_rows(db_session: Session, interview_order: ServiceOrder):
    zip_bytes = _zip_with_docx(
        [
            ("jane_doe_cv.docx", _docx_bytes("Jane Doe", phone="+44 7700 900111", email="jane@example.com")),
            ("john_smith.docx", _docx_bytes("John Smith", phone="+44 7700 900222", email="john@example.com")),
        ]
    )
    result = intake_cv_files(db_session, interview_order, [("cvs.zip", zip_bytes)])
    assert result["parsed_count"] == 2
    assert result["recipient_count"] == 2
    assert len(result["recipients"]) == 2
    names = {r["name"] for r in result["recipients"]}
    assert "Jane Doe" in names or "jane doe".title() in names
    assert result["summary"]["total"] == 2


def test_intake_mixed_zip_and_csv(db_session: Session, interview_order: ServiceOrder):
    import csv
    import io

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["name", "phone", "email"])
    writer.writerow(["Alice Example", "+447700900111", "alice@example.com"])
    csv_bytes = buf.getvalue().encode("utf-8")
    zip_bytes = _zip_with_docx([("bob_cv.docx", _docx_bytes("Bob Example"))])
    result = intake_mixed_files(
        db_session,
        interview_order,
        [("list.csv", csv_bytes), ("cvs.zip", zip_bytes)],
    )
    assert result["recipient_count"] >= 2
    assert len(result["recipients"]) >= 2


def test_empty_interview_draft_detection(db_session: Session, interview_order: ServiceOrder):
    assert is_empty_interview_draft(interview_order, recipient_count=0) is True

    interview_order.title = "Senior Engineer"
    assert is_empty_interview_draft(interview_order, recipient_count=0) is False

    interview_order.title = "Interview draft"
    interview_order.config_json = '{"role": "Engineer"}'
    assert is_empty_interview_draft(interview_order, recipient_count=0) is False


def test_purge_empty_interview_drafts(db_session: Session, interview_order: ServiceOrder):
    org_id = interview_order.org_id
    user_id = interview_order.user_id
    kept = create_new_interview_draft(db_session, org_id=org_id, user_id=user_id)
    assert db_session.get(ServiceOrder, interview_order.id) is None
    assert db_session.get(ServiceOrder, kept.id) is not None
    deleted = purge_empty_interview_drafts(db_session, org_id=org_id, keep_order_id=kept.id)
    assert deleted == 0


def test_abandon_empty_interview_draft(db_session: Session, interview_order: ServiceOrder):
    org_id = interview_order.org_id
    assert abandon_empty_interview_draft(db_session, org_id=org_id, order_id=interview_order.id) is True
    assert db_session.get(ServiceOrder, interview_order.id) is None
